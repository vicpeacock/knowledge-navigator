import os
import io
from pathlib import Path
from typing import Optional, Dict, Any, Tuple
import PyPDF2
from docx import Document
import openpyxl
from PIL import Image
# Note: python-magic not installed, using extension-based MIME type detection


class FileProcessor:
    """Service for processing uploaded files and extracting text content"""
    
    @staticmethod
    def extract_text(filepath: str, mime_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract text content from a file
        
        Returns:
            dict with 'text' and 'metadata' keys
        """
        path = Path(filepath)
        
        if not mime_type:
            # Guess MIME type from extension
            ext = path.suffix.lower()
            mime_types = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.xls': 'application/vnd.ms-excel',
                '.txt': 'text/plain',
                '.md': 'text/markdown',
            }
            mime_type = mime_types.get(ext, 'application/octet-stream')
        
        text = ""
        metadata: Dict[str, Any] = {
            "mime_type": mime_type,
            "file_size": path.stat().st_size,
        }
        
        try:
            if mime_type == "application/pdf":
                text, metadata = FileProcessor._extract_from_pdf(filepath)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                text, metadata = FileProcessor._extract_from_docx(filepath)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            ]:
                text, metadata = FileProcessor._extract_from_xlsx(filepath)
            elif mime_type.startswith("text/"):
                text, metadata = FileProcessor._extract_from_text(filepath)
            elif mime_type.startswith("image/"):
                # For images, we can't extract text directly
                # Would need OCR (like Tesseract) for this
                text = ""
                metadata["is_image"] = True
            else:
                text = ""
                metadata["unsupported"] = True
        except Exception as e:
            metadata["error"] = str(e)
            text = ""
        
        return {
            "text": text,
            "metadata": metadata,
        }
    
    @staticmethod
    def extract_text_from_bytes(
        file_bytes: bytes,
        mime_type: Optional[str],
        filename: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Extract text content from file bytes (for Cloud Storage).
        
        Args:
            file_bytes: File content as bytes
            mime_type: MIME type of the file
            filename: Optional filename (for extension detection if mime_type not provided)
        
        Returns:
            dict with 'text' and 'metadata' keys
        """
        # If no MIME type provided, try to guess from filename
        if not mime_type and filename:
            ext = Path(filename).suffix.lower()
            mime_types = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.xls': 'application/vnd.ms-excel',
                '.txt': 'text/plain',
                '.md': 'text/markdown',
            }
            mime_type = mime_types.get(ext, 'application/octet-stream')
        
        text = ""
        metadata: Dict[str, Any] = {
            "mime_type": mime_type or "application/octet-stream",
            "source": "bytes",
        }
        
        try:
            if mime_type == "application/pdf":
                text, metadata = FileProcessor._extract_from_pdf_bytes(file_bytes)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                text, metadata = FileProcessor._extract_from_docx_bytes(file_bytes)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            ]:
                text, metadata = FileProcessor._extract_from_xlsx_bytes(file_bytes)
            elif mime_type and mime_type.startswith("text/"):
                text, metadata = FileProcessor._extract_from_text_bytes(file_bytes)
            elif mime_type and mime_type.startswith("image/"):
                text = ""
                metadata["is_image"] = True
            else:
                text = ""
                metadata["unsupported"] = True
        except Exception as e:
            metadata["error"] = str(e)
            text = ""
        
        return {
            "text": text,
            "metadata": metadata,
        }
    
    @staticmethod
    def _extract_from_pdf_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF bytes"""
        text_parts = []
        metadata = {"pages": 0}
        
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        metadata["pages"] = len(pdf_reader.pages)
        
        for page in pdf_reader.pages:
            text_parts.append(page.extract_text())
        
        return "\n\n".join(text_parts), metadata
    
    @staticmethod
    def _extract_from_docx_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Word document bytes"""
        doc_file = io.BytesIO(file_bytes)
        doc = Document(doc_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        metadata = {"paragraphs": len(text_parts)}
        return "\n\n".join(text_parts), metadata
    
    @staticmethod
    def _extract_from_xlsx_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Excel file bytes"""
        excel_file = io.BytesIO(file_bytes)
        workbook = openpyxl.load_workbook(excel_file, data_only=True)
        
        text_parts = []
        sheet_names = []
        
        for sheet_name in workbook.sheetnames:
            sheet_names.append(sheet_name)
            sheet = workbook[sheet_name]
            
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)
        
        metadata = {"sheets": sheet_names, "rows": len(text_parts)}
        return "\n\n".join(text_parts), metadata
    
    @staticmethod
    def _extract_from_text_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from text file bytes"""
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode("latin-1")
            except UnicodeDecodeError:
                text = file_bytes.decode("utf-8", errors="ignore")
        
        metadata = {"encoding": "utf-8"}
        return text, metadata
    
    @staticmethod
    def _extract_from_pdf(filepath: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF"""
        text_parts = []
        metadata = {"pages": 0}
        
        with open(filepath, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata["pages"] = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
        
        return "\n\n".join(text_parts), metadata
    
    @staticmethod
    def _extract_from_docx(filepath: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Word document"""
        doc = Document(filepath)
        text_parts = [paragraph.text for paragraph in doc.paragraphs]
        
        metadata = {
            "paragraphs": len(doc.paragraphs),
        }
        
        return "\n\n".join(text_parts), metadata
    
    @staticmethod
    def _extract_from_xlsx(filepath: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Excel file"""
        workbook = openpyxl.load_workbook(filepath)
        text_parts = []
        
        metadata = {
            "sheets": workbook.sheetnames,
            "sheet_count": len(workbook.sheetnames),
        }
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_text = []
            
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) if cell else "" for cell in row)
                sheet_text.append(row_text)
            
            text_parts.append(f"Sheet: {sheet_name}\n" + "\n".join(sheet_text))
        
        return "\n\n".join(text_parts), metadata
    
    @staticmethod
    def _extract_from_text(filepath: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        with open(filepath, "r", encoding="utf-8", errors="ignore") as file:
            text = file.read()
        
        metadata = {
            "lines": len(text.splitlines()),
        }
        
        return text, metadata

